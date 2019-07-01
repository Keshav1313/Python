from docx import Document
from docx.shared import Inches


      ########HOW TO MAKE A WORD DOCUMENT##########
# document=Document()
# document.add_heading("This is Practice time. I very am having very less knowledge")
# document.add_paragraph("Working Hard will show its result very soon")
# document.save('Resume.docx')


# import docx2txt
# my_text = docx2txt.process("test.docx")
# # # print(my_text)

# import sys
# # import os
# # def file_read_from_tail(fname,lines):
# #         bufsize = 8192
# #         fsize = os.stat(fname).st_size
# #         iter = 0
# #         with open(fname) as f:
# #                 if bufsize > fsize:
# #                         bufsize = fsize-1
# #                         data = []
# #                         while True:
# #                                 iter +=1
# #                                 f.seek(fsize-bufsize*iter)
# #                                 data.extend(f.readlines())
# #                                 if len(data) >= lines or f.tell() == 0:
# #                                         print(''.join(data[-lines:]))
# #                                         break
# #
# # file_read_from_tail('hello.txt',2)


#BEST CODE FOR FINDING LAST LINE/BYTES OF ANY TEXT FILE

# def read_last_lines(filename, lines):
#     with open(filename, "r") as text_file:
#         contents = text_file.readlines()[-lines:]
#         for line in contents:
#             print(line, end="")
#
# read_last_lines("hello.txt", 11)

# name='keshav_sinha'
# print(name[-5:])

# import random
# while True:
#     user_number=input("Enter your number : \n")
#     user_number = int(user_number)
#     winning_number=random.randint(0,10)
#     if user_number == winning_number:
#         print("You guessed correct number.")
#         break
#     else:
#         if user_number < winning_number:
#             print("Too Low")
#         else:
#             print("Too High")
#             print("Try again")


# total=0
# i=1
# while i<=10:
#     total=total+i
#     i=i+1
# print(total)


#PROGRAM TO PRINT THE TOTAL OF GIVEN NUMBER##########
# sum=0
# i=1
# user_num=input("Enter a number to get its total:")
# user_num=int(user_num)
# while i<=user_num:
#     sum=sum+i
#     i+=1
# print("Sum of your given number is :"+str(sum))

#PROGRAM TO PRINT THE SUM OF MORE THAN 2 DIGIT NUMBER

# num=input("Enter a digit more than 2 :")
# sum=0
# i=0
# while i<len(num):
#     sum+=int(num[i])
#     i+=1
# print(sum)

#PROGRAM TO PRINT THE NUMBER OF TIMES AN ALPHABET IN NAME
# name=input("Enter your name or any string:")
# print(name)
# temp_var=""
# i=0
# while i<len(name):
#     if name[i] not in temp_var:
#         temp_var+=name[i]
#         print( name[i] + " : " + str(name.count(name[i])))
#     i+=1

#FOR  LOOP
# total=0
# num=input("enter number:")
# for i in range(0,len(num)):
#     total+=int(num[i])
# print(total)


#PROGRAM TO PRINT THE NUMBER OF TIMES AN ALPHABET IN STRING
# name=input("enter a name:")
# temp=""
# for i in range(0,len(name)):
#     if name[i] not in temp:
#         print(name[i] + ":" + str(name.count(name[i])))
#         temp+=temp+name[i]



import os
import sys
fname=input("Enter the file name:")
if os.path.isfile(fname):
    print("File exist:",fname)
    f=open(fname,'r')
else:
    print("File doesn't exist:")
    sys.exit(0)
lcount=wcount=ccount=0
for line in f:
    lcount=lcount+1
    ccount=ccount+len(line)
    words=line.split()
    wcount=wcount+len(words)

print("The number of lines :",lcount)
print("The number of word count:",wcount)
print("The number of characters:",ccount)
print(os.getcwd())






































